import html
import re
from collections.abc import Iterator
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

FIELD_PATTERN = re.compile(r"\{\{\s*([^{}\n\r]+?)\s*\}\}")
DOCX_XML_PARTS_PATTERN = re.compile(r"word/(document|header\d+|footer\d+)\.xml$")


def _iter_paragraphs_from_parent(parent) -> Iterator[Paragraph]:
    if isinstance(parent, Table):
        for row in parent.rows:
            for cell in row.cells:
                yield from _iter_paragraphs_from_parent(cell)
        return

    if hasattr(parent, "paragraphs") and hasattr(parent, "tables"):
        for paragraph in parent.paragraphs:
            yield paragraph
        for table in parent.tables:
            yield from _iter_paragraphs_from_parent(table)
        return


def _iter_document_paragraphs(document) -> Iterator[Paragraph]:
    yield from _iter_paragraphs_from_parent(document)

    seen_header_footer_parts: set[int] = set()
    for section in document.sections:
        for part in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            part_id = id(part._element)
            if part_id in seen_header_footer_parts:
                continue

            seen_header_footer_parts.add(part_id)
            yield from _iter_paragraphs_from_parent(part)


def _paragraphs_text(document) -> list[str]:
    return [paragraph.text.strip() for paragraph in _iter_document_paragraphs(document) if paragraph.text.strip()]


def _extract_fields_from_text(text: str) -> list[str]:
    fields: list[str] = []

    for match in FIELD_PATTERN.findall(text):
        field = match.strip()
        if field and field not in fields:
            fields.append(field)

    return fields


def _extract_docx_xml_text(doc_path: str | Path) -> str:
    fragments: list[str] = []

    with ZipFile(doc_path) as archive:
        for member in archive.namelist():
            if not DOCX_XML_PARTS_PATTERN.fullmatch(member):
                continue

            xml = archive.read(member).decode("utf-8", errors="ignore")
            xml = re.sub(r"<w:tab[^>]*/>", "\t", xml)
            xml = re.sub(r"<w:(?:br|cr)[^>]*/>", "\n", xml)
            xml = re.sub(r"</w:p>", "\n", xml)
            xml = re.sub(r"</w:tr>", "\n", xml)
            xml = re.sub(r"</w:tc>", "\t", xml)
            xml = re.sub(r"<[^>]+>", "", xml)
            text = html.unescape(xml)

            normalized_lines = [
                re.sub(r"[ \t]+", " ", line).strip()
                for line in text.splitlines()
            ]
            normalized_text = "\n".join(line for line in normalized_lines if line)
            if normalized_text:
                fragments.append(normalized_text)

    return "\n".join(fragments)


def inspect_template(doc_path: str | Path) -> dict[str, object]:
    document = Document(str(doc_path))
    text = "\n".join(_paragraphs_text(document))

    fields: list[str] = []
    for source in (text, _extract_docx_xml_text(doc_path)):
        for field in _extract_fields_from_text(source):
            if field not in fields:
                fields.append(field)

    return {"text": text, "detected_fields": fields}


def extract_template_text(doc_path: str | Path) -> str:
    return str(inspect_template(doc_path)["text"])


def extract_fields(doc_path: str | Path) -> list[str]:
    return list(inspect_template(doc_path)["detected_fields"])


def _replace_in_paragraph(paragraph: Paragraph, data: dict[str, object]) -> None:
    text = paragraph.text
    replaced = FIELD_PATTERN.sub(lambda match: "" if data.get(match.group(1).strip()) is None else str(data.get(match.group(1).strip())), text)

    if replaced != text:
        paragraph.text = replaced


def render_docx(template_path: str | Path, data: dict[str, object]) -> bytes:
    document = Document(str(template_path))

    for paragraph in _iter_document_paragraphs(document):
        _replace_in_paragraph(paragraph, data)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def has_markers(doc_path: str | Path) -> bool:
    return bool(extract_fields(doc_path))


def extract_marked_values(doc_path: str | Path, expected_fields: list[str] | None = None) -> dict[str, str | None]:
    text = extract_template_text(doc_path)
    data: dict[str, str | None] = {field: None for field in expected_fields or []}

    for line in text.splitlines():
        matches = list(FIELD_PATTERN.finditer(line))
        if not matches:
            continue

        for match in matches:
            field = match.group(1).strip()
            suffix = line[match.end():].strip(" :-\t")
            prefix = line[: match.start()].strip()

            value = suffix or None
            if not value and ":" in prefix:
                value = prefix.split(":", 1)[-1].strip() or None

            data[field] = value

    return data
